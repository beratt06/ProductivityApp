import sys
import os
import re
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """Sets the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding/margins in dxas (1/20 of a pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_border(paragraph, position='bottom', color='CCCCCC', size='12'):
    """Adds a border to a paragraph (useful for horizontal rules)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bdr = OxmlElement(f'w:{position}')
    bdr.set(qn('w:val'), 'single')
    bdr.set(qn('w:sz'), size)
    bdr.set(qn('w:space'), '4')
    bdr.set(qn('w:color'), color)
    pBdr.append(bdr)
    pPr.append(pBdr)

def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # Configure page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style definitions
    styles = doc.styles
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    in_code_block = False
    code_text = []
    
    in_table = False
    table_rows = []
    
    in_mermaid = False
    mermaid_text = []

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Handle Code Blocks
        if stripped.startswith("```"):
            if in_code_block or in_mermaid:
                # Ending block
                if in_mermaid:
                    p = doc.add_paragraph()
                    add_border(p, 'left', '008000', '24')
                    run = p.add_run("[Diyagram: Mermaid Kodu Aşağıdadır. Çizim için çevrimiçi render edebilirsiniz]\n")
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                    p.add_run("".join(mermaid_text))
                    p.paragraph_format.left_indent = Inches(0.25)
                    in_mermaid = False
                    mermaid_text = []
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    # Format as code block style
                    p_run = p.add_run("".join(code_text))
                    p_run.font.name = 'Consolas'
                    p_run.font.size = Pt(9.5)
                    p_run.font.color.rgb = RGBColor(0x00, 0x66, 0x99)
                    
                    # Set background
                    # Note: Full paragraph background in python-docx is hard, so we just format font.
                    in_code_block = False
                    code_text = []
            else:
                # Starting block
                if "mermaid" in stripped:
                    in_mermaid = True
                else:
                    in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_text.append(line)
            i += 1
            continue
            
        if in_mermaid:
            mermaid_text.append(line)
            i += 1
            continue

        # Handle Tables
        if stripped.startswith("|"):
            in_table = True
            table_rows.append(stripped)
            i += 1
            continue
        elif in_table:
            # Table ended, render it
            render_table(doc, table_rows)
            table_rows = []
            in_table = False

        # Handle horizontal rules
        if stripped == "---":
            p = doc.add_paragraph()
            add_border(p, 'bottom', 'CCCCCC', '6')
            i += 1
            continue

        # Handle Alert Boxes (e.g. > [!IMPORTANT] ...)
        if stripped.startswith("> [!"):
            alert_type = re.search(r'\[!(.*)\]', stripped)
            alert_name = alert_type.group(1) if alert_type else "NOTE"
            i += 1
            alert_content = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                alert_content.append(lines[i].strip().lstrip("> ").strip())
                i += 1
            
            p = doc.add_paragraph()
            add_border(p, 'left', 'D9534F' if alert_name in ['IMPORTANT', 'WARNING', 'CAUTION'] else '5BC0DE', '24')
            p.paragraph_format.left_indent = Inches(0.25)
            
            run_title = p.add_run(f"[{alert_name}] ")
            run_title.bold = True
            if alert_name in ['IMPORTANT', 'WARNING', 'CAUTION']:
                run_title.font.color.rgb = RGBColor(0xD9, 0x53, 0x4F)
            else:
                run_title.font.color.rgb = RGBColor(0x5B, 0xC0, 0xDE)
                
            p.add_run(" ".join(alert_content))
            continue

        # Handle Headings
        if stripped.startswith("# "):
            p = doc.add_heading(level=1)
            run = p.add_run(stripped[2:])
            run.font.name = 'Calibri'
            run.font.size = Pt(20)
            run.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Navy Blue
            # Add bottom border
            add_border(p, 'bottom', '1B365D', '18')
            i += 1
            continue
        elif stripped.startswith("## "):
            p = doc.add_heading(level=2)
            run = p.add_run(stripped[3:])
            run.font.name = 'Calibri'
            run.font.size = Pt(15)
            run.bold = True
            run.font.color.rgb = RGBColor(0x2E, 0x5B, 0x88) # Slate Blue
            i += 1
            continue
        elif stripped.startswith("### "):
            p = doc.add_heading(level=3)
            run = p.add_run(stripped[4:])
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(0x4A, 0x77, 0x9D)
            i += 1
            continue

        # Handle Bullet Points
        if stripped.startswith("* ") or stripped.startswith("- "):
            bullet_text = stripped[2:]
            # Clean markdown bold syntax like **text**
            bullet_text = clean_markdown_bold(bullet_text)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            
            # Format markdown links [text](url) -> text
            parse_and_add_text_with_links(p, bullet_text)
            i += 1
            continue

        # Handle Normal Paragraphs
        if stripped:
            stripped = clean_markdown_bold(stripped)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            parse_and_add_text_with_links(p, stripped)
            
        i += 1

    # Save final doc
    doc.save(docx_path)
    print(f"Successfully converted MD to DOCX: {docx_path}")

def clean_markdown_bold(text):
    # This is a basic cleaner for markdown bold
    # Real parser would support inline bolding in Word
    return text

def parse_and_add_text_with_links(paragraph, text):
    # Quick parser for inline markdown formatting: **bold** and [text](link)
    # Let's break text by bold markers first
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            sub_text = part[2:-2]
            run = paragraph.add_run(sub_text)
            run.bold = True
        else:
            # Parse links in non-bold text
            sub_parts = re.split(r'(\[.*?\]\(.*?\))', part)
            for sub_part in sub_parts:
                link_match = re.match(r'\[(.*?)\]\((.*?)\)', sub_part)
                if link_match:
                    link_text = link_match.group(1)
                    run = paragraph.add_run(link_text)
                    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
                    run.underline = True
                else:
                    paragraph.add_run(sub_part)

def render_table(doc, rows):
    # Split each row by '|' and strip spaces
    table_data = []
    for r in rows:
        cells = [c.strip() for c in r.split("|")[1:-1]]
        # Skip divider row (e.g. | :--- | :--- |)
        if all(re.match(r'^:?-+:?$', c) for c in cells if c):
            continue
        table_data.append(cells)
        
    if not table_data:
        return
        
    num_cols = len(table_data[0])
    num_rows = len(table_data)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.autofit = True
    
    for r_idx, row in enumerate(table_data):
        for c_idx, val in enumerate(row):
            # Clean markdown bold/links in table cell value
            cell = table.cell(r_idx, c_idx)
            cell.text = "" # Clear default text
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            
            # Format text in cell
            parse_and_add_text_with_links(p, val)
            
            # Style header row differently
            if r_idx == 0:
                set_cell_background(cell, "1B365D") # Navy Blue Header
                # Set text color to white and bold
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                # Zebra striping
                if r_idx % 2 == 0:
                    set_cell_background(cell, "F2F5F8") # Light grey-blue
                else:
                    set_cell_background(cell, "FFFFFF")
                    
            # Padding
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)

if __name__ == "__main__":
    if len(sys.argv) < 3:
        md_file = "proje_raporu_taslagi.md"
        docx_file = "proje_raporu.docx"
    else:
        md_file = sys.argv[1]
        docx_file = sys.argv[2]
        
    convert_md_to_docx(md_file, docx_file)
